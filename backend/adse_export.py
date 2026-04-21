"""
ADSE - Módulo de Exportación a Word (.docx)
Genera documentos con formato oficial SEP: membrete, estructura, firma y sello.
Adaptado para directores de Jardín de Niños (secundaria).
"""

import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


def crear_documento_sep(contenido, metadata=None):
    """
    Genera un documento Word con formato oficial SEP.

    Args:
        contenido: Texto del documento generado por la IA
        metadata: Dict con datos adicionales (funcion, herramienta, titulo,
                  escuela, director, fecha)

    Returns:
        BytesIO con el documento Word listo para descargar
    """
    if metadata is None:
        metadata = {}

    doc = Document()

    # ============================================
    # CONFIGURAR MÁRGENES
    # ============================================
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================
    # ESTILOS
    # ============================================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ============================================
    # MEMBRETE SEP
    # ============================================
    membrete = doc.add_paragraph()
    membrete.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = membrete.add_run('SECRETARÍA DE EDUCACIÓN PÚBLICA')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)

    submembrete = doc.add_paragraph()
    submembrete.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = submembrete.add_run('SUBSECRETARÍA DE EDUCACIÓN BÁSICA')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)

    # Escuela
    escuela = metadata.get('escuela', '')
    if escuela:
        esc_p = doc.add_paragraph()
        esc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = esc_p.add_run(f'Jardín de Niños {escuela}')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x9C, 0x27, 0xB0)

    # Línea separadora
    linea = doc.add_paragraph()
    linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linea.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
    linea.paragraph_format.space_after = Pt(12)

    # ============================================
    # DATOS DEL DOCUMENTO
    # ============================================
    fecha = metadata.get('fecha', datetime.now().strftime('%d de %B de %Y'))
    funcion = metadata.get('funcion', '')
    herramienta = metadata.get('herramienta', '')
    titulo = metadata.get('titulo', 'Documento ADSE')

    # Fecha alineada a la derecha
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fecha_p.add_run(f'Fecha: {fecha}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    # Obligación y herramienta
    if funcion:
        func_p = doc.add_paragraph()
        run = func_p.add_run(f'Obligación: ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = func_p.add_run(funcion)
        run2.font.size = Pt(10)

    if herramienta:
        herr_p = doc.add_paragraph()
        run = herr_p.add_run(f'Herramienta: ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = herr_p.add_run(herramienta)
        run2.font.size = Pt(10)

    # Espacio
    doc.add_paragraph('')

    # ============================================
    # TÍTULO DEL DOCUMENTO
    # ============================================
    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_p.add_run(titulo.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
    titulo_p.paragraph_format.space_after = Pt(16)

    # ============================================
    # CONTENIDO DEL DOCUMENTO
    # ============================================
    lineas = contenido.split('\n')

    for linea_texto in lineas:
        linea_strip = linea_texto.strip()

        if not linea_strip:
            doc.add_paragraph('')
            continue

        # Detectar encabezados
        if linea_strip.startswith('###'):
            p = doc.add_paragraph()
            run = p.add_run(linea_strip.replace('###', '').strip())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
        elif linea_strip.startswith('##'):
            p = doc.add_paragraph()
            run = p.add_run(linea_strip.replace('##', '').strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
        elif linea_strip.startswith('#'):
            p = doc.add_paragraph()
            run = p.add_run(linea_strip.replace('#', '').strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
        elif linea_strip.startswith('---') or linea_strip.startswith('==='):
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
        elif linea_strip.startswith('- ') or linea_strip.startswith('• '):
            texto = linea_strip.lstrip('-•').strip()
            p = doc.add_paragraph(texto, style='List Bullet')
        elif linea_strip[0:3].replace('.', '').replace(')', '').isdigit():
            p = doc.add_paragraph(linea_strip, style='List Number')
        elif linea_strip.isupper() and len(linea_strip) > 3:
            p = doc.add_paragraph()
            run = p.add_run(linea_strip)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
            p.paragraph_format.space_before = Pt(12)
        elif linea_strip.endswith(':') and len(linea_strip) < 60:
            p = doc.add_paragraph()
            run = p.add_run(linea_strip)
            run.bold = True
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(linea_strip)

    # ============================================
    # ESPACIO PARA FIRMA
    # ============================================
    doc.add_paragraph('')
    doc.add_paragraph('')

    firma_sep = doc.add_paragraph()
    firma_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = firma_sep.add_run('━' * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)

    atentamente = doc.add_paragraph()
    atentamente.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = atentamente.add_run('Atentamente')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    linea_firma = doc.add_paragraph()
    linea_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linea_firma.add_run('_' * 40)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    director = metadata.get('director', metadata.get('supervisor', 'Nombre del Director(a)'))
    nombre_p = doc.add_paragraph()
    nombre_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = nombre_p.add_run(director)
    run.bold = True
    run.font.size = Pt(11)

    cargo_p = doc.add_paragraph()
    cargo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cargo_p.add_run('Director(a) del Jardín de Niños')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    # Espacio para sello
    doc.add_paragraph('')
    sello_p = doc.add_paragraph()
    sello_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sello_p.add_run('[Sello Oficial]')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xb0, 0xb0, 0xb0)
    run.italic = True

    # ============================================
    # PIE DE PÁGINA
    # ============================================
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.paragraph_format.space_before = Pt(20)
    run = pie.add_run('Documento generado por ADSE — Agente Director Secundaria de Educación')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xb0, 0xb0, 0xb0)
    run.italic = True

    # ============================================
    # GUARDAR EN MEMORIA
    # ============================================
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def generar_nombre_archivo(funcion, herramienta, extension='docx'):
    """Genera nombre de archivo limpio para descarga."""
    fecha = datetime.now().strftime('%Y%m%d')
    nombre = f"{funcion}_{herramienta}_{fecha}"
    nombre = ''.join(c if c.isalnum() or c in ('_', '-') else '_' for c in nombre)
    nombre = nombre[:60]
    return f"{nombre}.{extension}"
